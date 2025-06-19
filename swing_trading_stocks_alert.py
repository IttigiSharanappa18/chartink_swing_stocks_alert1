import time
import pandas as pd
import asyncio
from docx import Document
from telegram import Bot
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.options import Options

# Replace these with your actual credentials and details
username = 'ittigisharan25@gmail.com'
password = '13254768'
telegram_token = '7629549320:AAEVpmRMhASfwCktIBd_e19BlYZ5xH19JfI'
telegram_chat_id = '-1002259128052'

print("Starting script...")

# Login and get stocks function
def get_stocks():
    Final_stock_names = []
    scan_links_dict = {
        'Bearish-15-Min-Breakdown' : 'https://chartink.com/screener/copy-bearish-swing-trading-srimantha-15-mins-breakdown-312',
        'Bearish-1-Hour-Breakdown' : 'https://chartink.com/screener/copy-bearish-swing-trading-srimantha-1-hour-breakdown-360',
        'Bullish-15-Min-Breakout' : 'https://chartink.com/screener/copy-bullish-swing-trading-srimantha-15-mins-breakout-314',
        'Bullish-1-Hour-Breakout' : 'https://chartink.com/screener/copy-bullish-swing-trading-srimantha-hourly-breakout-344',
        'Bullish-Daily-Breakout':'https://chartink.com/screener/copy-bullish-swing-trading-srimantha-daily-breakout-373',
        #'Reversal-Daily-Positional':'https://chartink.com/screener/copy-reversal-daily-positional-scanner-330',
        #'Reversal-1-Hour':'https://chartink.com/screener/copy-reversal-1-hour-scanner-304',
        #'Reversal-5mins':'https://chartink.com/screener/copy-reversal-5-mins-scanner-310',
        #'EMA-Swing-stocks' :'https://chartink.com/screener/sharan-ittigi'
    }
    Final_stock_names_dict = dict()

    for scan_link_name,scan_links in scan_links_dict.items():
        print("Initializing WebDriver...")
        #service = Service(r'.github/workflows/chromedriver.exe')
        #driver = webdriver.Chrome(service=service)
        

        options = Options()
        options.add_argument('--headless')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        driver.get('https://www.chartink.com/login')
        print("Logging in...")
        WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.NAME, 'email')))
        driver.find_element(By.NAME, 'email').send_keys(username)
        driver.find_element(By.NAME, 'password').send_keys(password)
        driver.find_element(By.CLASS_NAME, 'g-recaptcha').send_keys(Keys.RETURN)
        print("Navigating to SCAN page...")
        driver.get(scan_links)
        driver.find_element(By.CLASS_NAME, 'btn.btn-success.run_scan_button').click()
        #class="btn btn-success run_scan_button"
        print("Waiting for results...")
        time.sleep(5)
        stocks = driver.find_elements(By.CLASS_NAME, 'text-teal-700')  # Adjust the class name based on actual HTML
        print("Extracting stock names...")
        stock_names = [stock.text for stock in stocks]
        stock_names = stock_names[::4]
        for each in stock_names:
            Final_stock_names.append(each)
        Final_stock_names_dict[scan_link_name] = str(stock_names)
        driver.quit()
        print("Stocks retrieved:", Final_stock_names_dict)
    return Final_stock_names_dict

# Send stocks to Telegram function
async def send_to_telegram(stocks):
    bot = Bot(token=telegram_token)
    for swing_type,stock_list in stocks.items():
        if len(stock_list) != 0:
            await bot.send_message(chat_id=telegram_chat_id, text= "#" +swing_type+ "#" + '\n' +str(stock_list))
        else:
            pass
    print("Stocks sent to Telegram")

# Save stocks to Excel function
async def save_to_file(stocks):
    # Create a Word document
    doc = Document()
    # Add a title
    doc.add_heading('Swing trading stocks', level=1)
    # Add dictionary items to the document
    for key, value in stocks.items():
        doc.add_heading(key, level=2)  # Add a heading for each key
        doc.add_paragraph(str(value))  # Add the corresponding list as a paragraph

    # Save the document
    doc.save('Swing_trading_stocks.docx')
    #df.to_('stocks.xlsx', index=False)
    print("Stocks saved to file")
    #word_file_path = "D:/Automotion_chartink/Swing_trading_stocks.docx"
    #bot = Bot(token=telegram_token)
    #await bot.send_document(chat_id=telegram_chat_id, document=open(word_file_path, "rb"))
    #print("Stocks file sent to telegram")

# Main function to run the whole process
def job():
    print("Job started...")
    stocks = get_stocks()
    asyncio.run(send_to_telegram(stocks))
    asyncio.run(save_to_file(stocks))
    print("Job completed")

# Schedule the job every 15 minutes during Indian market hours
#schedule.every(1).minutes.do(job)
print("Script scheduled to run every 1 minutes")
job()
#while True:
#    schedule.run_pending()
#    time.sleep(1)
